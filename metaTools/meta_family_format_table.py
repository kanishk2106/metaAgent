from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os
import re
from typing import Optional, Dict, Any, List, Union
def register(server):
    @server.tool()
    def format_table_with_optional_elements(
        filename: str,
        table_index: int,
        has_header_row: bool = False,
        border_style: Optional[str] = None,
        shading: Optional[List[Dict[str, Union[int, str]]]] = None,
        add_heading_text: Optional[str] = None,
        add_heading_level: int = 1,
        add_paragraph_text: Optional[str] = None,
        add_paragraph_style: Optional[str] = None,
        add_table_rows: Optional[int] = None,
        add_table_cols: Optional[int] = None,
        add_table_data: Optional[List[List[Any]]] = None,
        second_format_table_index: Optional[int] = None,
        second_format_table_has_header_row: bool = False,
        second_format_table_border_style: Optional[str] = None,
        second_format_table_shading: Optional[List[Dict[str, Union[int, str]]]] = None,
    ) -> Dict[str, Any]:
        operations_summary = {
            "mandatory_table_format": {"applied": False},
            "heading_added": {"applied": False},
            "paragraph_added": {"applied": False},
            "table_added": {"applied": False},
            "optional_table_format": {"applied": False},
        }
        try:
            filename = os.path.abspath(filename)
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith(".docx"):
                return {"error": "File must be a .docx file"}
            doc = Document(filename)
            total_tables = len(doc.tables)
            if not isinstance(table_index, int) or table_index < 0 or table_index >= total_tables:
                return {
                    "error": f"Mandatory table_index {table_index} out of range (0-{total_tables-1})"
                }
            allowed_borders = ["single", "double", "none", None]
            if border_style not in allowed_borders:
                return {
                    "error": f"Unsupported border_style '{border_style}'. Allowed: {allowed_borders}"
                }
            hex_pattern = re.compile(r"^[0-9A-Fa-f]{6}$")
            if add_heading_text is not None and (
                not isinstance(add_heading_level, int) or add_heading_level < 1 or add_heading_level > 9
            ):
                return {"error": "add_heading_level must be an integer between 1 and 9"}
            if (add_table_rows is not None and add_table_cols is None) or \
               (add_table_cols is not None and add_table_rows is None):
                return {"error": "Both add_table_rows and add_table_cols must be provided to add a table."}
            if add_table_rows is not None:
                if not isinstance(add_table_rows, int) or not isinstance(add_table_cols, int):
                    return {"error": "add_table_rows and add_table_cols must be integers"}
                if add_table_rows < 1 or add_table_cols < 1:
                    return {"error": "add_table_rows and add_table_cols must be >= 1"}
                if add_table_data is not None:
                    if not isinstance(add_table_data, list):
                        return {"error": "add_table_data must be a 2D list"}
                    for r_data in add_table_data:
                        if not isinstance(r_data, list):
                            return {"error": "Each row in add_table_data must be a list"}
            if second_format_table_index is not None:
                if not isinstance(second_format_table_index, int) or \
                   second_format_table_index < 0 or second_format_table_index >= total_tables:
                    return {
                        "error": f"Optional second_format_table_index {second_format_table_index} "
                                 f"out of range (0-{total_tables-1})"
                    }
                if second_format_table_border_style not in allowed_borders:
                    return {
                        "error": f"Unsupported second_format_table_border_style '{second_format_table_border_style}'. "
                                 f"Allowed: {allowed_borders}"
                    }
            current_table = doc.tables[table_index]
            if has_header_row and len(current_table.rows) > 0:
                try:
                    current_table.style = "Table Grid"
                except Exception:
                    pass  # Ignore if style not found, don't crash
            shaded_rows_mandatory = 0
            if shading:
                if not isinstance(shading, list):
                    return {"error": "shading must be a list of dict entries"}
                for entry in shading:
                    if not isinstance(entry, dict): continue
                    row_idx = entry.get("row")
                    color = entry.get("color", "FFFFFF")
                    if not isinstance(row_idx, int) or row_idx < 0 or row_idx >= len(current_table.rows): continue
                    if not isinstance(color, str) or not hex_pattern.match(color): continue
                    try:
                        for cell in current_table.rows[row_idx].cells:
                            shade = OxmlElement("w:shd")
                            shade.set(qn("w:fill"), color)
                            cell._element.get_or_add_tcPr().append(shade)
                        shaded_rows_mandatory += 1
                    except Exception: continue
            operations_summary["mandatory_table_format"] = {
                "applied": True,
                "table_index": table_index,
                "has_header_row": has_header_row,
                "border_style": border_style,
                "shading_rows": shaded_rows_mandatory,
                "message": f"Formatted table {table_index}"
            }
            if add_heading_text is not None:
                doc.add_heading(add_heading_text, level=add_heading_level)
                operations_summary["heading_added"] = {
                    "applied": True,
                    "text": add_heading_text,
                    "level": add_heading_level,
                    "message": f"Added heading '{add_heading_text}' at level {add_heading_level}"
                }
            if add_paragraph_text is not None:
                paragraph = doc.add_paragraph(add_paragraph_text)
                actual_style = None
                if add_paragraph_style:
                    try:
                        paragraph.style = add_paragraph_style
                        actual_style = add_paragraph_style
                    except KeyError:
                        return {
                            "error": f"Style '{add_paragraph_style}' not found in document for optional paragraph",
                            "filename": filename,
                        }
                operations_summary["paragraph_added"] = {
                    "applied": True,
                    "text_length": len(add_paragraph_text),
                    "style": actual_style or "default",
                    "message": f"Added paragraph with {len(add_paragraph_text)} characters"
                }
            if add_table_rows is not None and add_table_cols is not None:
                new_table = doc.add_table(rows=add_table_rows, cols=add_table_cols)
                populated_rows_new_table = 0
                if add_table_data:
                    for i, row_data in enumerate(add_table_data):
                        if i >= add_table_rows: break
                        for j, cell_value in enumerate(row_data):
                            if j >= add_table_cols: break
                            new_table.rows[i].cells[j].text = str(cell_value)
                        populated_rows_new_table += 1
                operations_summary["table_added"] = {
                    "applied": True,
                    "rows": add_table_rows,
                    "cols": add_table_cols,
                    "populated_rows": populated_rows_new_table,
                    "message": f"Added {add_table_rows}x{add_table_cols} table"
                }
            if second_format_table_index is not None:
                current_table = doc.tables[second_format_table_index]
                if second_format_table_has_header_row and len(current_table.rows) > 0:
                    try:
                        current_table.style = "Table Grid"
                    except Exception:
                        pass
                shaded_rows_optional = 0
                if second_format_table_shading:
                    if not isinstance(second_format_table_shading, list):
                        return {"error": "second_format_table_shading must be a list of dict entries"}
                    for entry in second_format_table_shading:
                        if not isinstance(entry, dict): continue
                        row_idx = entry.get("row")
                        color = entry.get("color", "FFFFFF")
                        if not isinstance(row_idx, int) or row_idx < 0 or row_idx >= len(current_table.rows): continue
                        if not isinstance(color, str) or not hex_pattern.match(color): continue
                        try:
                            for cell in current_table.rows[row_idx].cells:
                                shade = OxmlElement("w:shd")
                                shade.set(qn("w:fill"), color)
                                cell._element.get_or_add_tcPr().append(shade)
                            shaded_rows_optional += 1
                        except Exception: continue
                operations_summary["optional_table_format"] = {
                    "applied": True,
                    "table_index": second_format_table_index,
                    "has_header_row": second_format_table_has_header_row,
                    "border_style": second_format_table_border_style,
                    "shading_rows": shaded_rows_optional,
                    "message": f"Formatted second table {second_format_table_index}"
                }
            doc.save(filename)
            overall_message = "Formatted primary table. "
            if operations_summary["heading_added"]["applied"]:
                overall_message += "Added heading. "
            if operations_summary["paragraph_added"]["applied"]:
                overall_message += "Added paragraph. "
            if operations_summary["table_added"]["applied"]:
                overall_message += "Added new table. "
            if operations_summary["optional_table_format"]["applied"]:
                overall_message += "Formatted second table. "
            overall_message = overall_message.strip()
            return {
                "success": True,
                "filename": filename,
                "operations": operations_summary,
                "message": overall_message
            }
        except Exception as e:
            return {"error": f"Failed to perform document operations: {str(e)}"}