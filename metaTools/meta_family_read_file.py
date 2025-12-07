"""
MCP Tool: read_file_meta
Reads a file and optionally performs other file system or document operations.
"""

from docx import Document
import os
import subprocess
import shlex
from typing import Optional, Dict, Any, List


# Helper functions encapsulating base tool logic
# These are kept outside the register function for better organization and testability.

def _read_file_logic(
    path: str,
    start_line: Optional[int] = None,
    end_line: Optional[int] = None,
    max_bytes: int = 200_000
) -> Dict[str, Any]:
    """Reads the contents of a text file."""
    try:
        path = os.path.abspath(os.path.expanduser(path))

        if not os.path.exists(path):
            return {"error": f"File does not exist: {path}"}

        if not os.path.isfile(path):
            return {"error": f"Path is not a file: {path}"}

        file_size = os.path.getsize(path)
        if file_size > max_bytes:
            return {
                "error": f"File too large ({file_size} bytes). "
                         f"Max allowed is {max_bytes} bytes."
            }

        with open(path, "r", encoding="utf-8") as f:
            if start_line is not None or end_line is not None:
                if start_line is not None and start_line < 1:
                    return {"error": "start_line must be >= 1"}
                if end_line is not None and end_line < 1:
                    return {"error": "end_line must be >= 1"}

                lines = f.readlines()
                start = (start_line - 1) if start_line else 0
                end = end_line if end_line else len(lines)

                if start >= len(lines):
                    return {"error": "start_line beyond file length"}
                if end < start:
                    return {"error": "end_line must be >= start_line"}

                content = "".join(lines[start:end])
            else:
                content = f.read()

        return {
            "success": True,
            "path": path,
            "size": len(content),
            "content": content,
            "message": f"Read {len(content)} bytes from {path}"
        }

    except UnicodeDecodeError:
        return {"error": f"File is not UTF-8 text: {path}"}
    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": f"Error reading file: {str(e)}"}


def _add_paragraph_logic(filename: str, text: str, style: Optional[str] = None) -> Dict[str, Any]:
    """Adds a paragraph to a Word document."""
    try:
        if not os.path.exists(filename):
            return {"error": f"File not found: {filename}"}
        if not filename.endswith('.docx'):
            return {"error": "File must be a .docx file"}

        doc = Document(filename)
        paragraph = doc.add_paragraph(text)

        if style:
            try:
                paragraph.style = style
            except KeyError:
                return {"error": f"Style '{style}' not found in document"}

        doc.save(filename)

        return {
            "success": True,
            "filename": filename,
            "text": text,
            "style": style if style else "default",
            "message": f"Added paragraph with {len(text)} characters"
        }

    except Exception as e:
        return {"error": f"Failed to add paragraph: {str(e)}"}


def _execute_command_logic(command: str, timeout: int = 10) -> Dict[str, Any]:
    """Safely executes a whitelisted shell command."""
    allowed_commands = ["ls", "pwd", "whoami", "echo", "date"]

    try:
        parts = shlex.split(command)
        if not parts:
            return {"error": "Empty command"}
        cmd = parts[0]

        if cmd not in allowed_commands:
            return {
                "error": f"Command '{cmd}' is not allowed.",
                "allowed": allowed_commands
            }

        result = subprocess.run(
            parts,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=timeout
        )

        return {
            "success": True,
            "command": command,
            "stdout": result.stdout,
            "stderr": result.stderr,
            "returncode": result.returncode,
            "message": f"Command '{cmd}' executed safely"
        }

    except subprocess.TimeoutExpired:
        return {"error": f"Command timed out after {timeout} seconds"}
    except Exception as e:
        return {"error": f"Execution failed: {str(e)}"}


def _list_directory_logic(path: str = ".") -> Dict[str, Any]:
    """Lists files and subdirectories in the specified directory."""
    try:
        path = os.path.abspath(os.path.expanduser(path))

        if not os.path.exists(path):
            return {"error": f"Path does not exist: {path}"}
        if not os.path.isdir(path):
            return {"error": f"Path is not a directory: {path}"}

        items = os.listdir(path)
        detailed_items = []
        for item in items:
            item_path = os.path.join(path, item)
            is_dir = os.path.isdir(item_path)
            is_link = os.path.islink(item_path)

            detailed_items.append({
                "name": item,
                "type": "directory" if is_dir else "file",
                "absolute_path": os.path.abspath(item_path),
                "is_symlink": is_link
            })

        return {
            "success": True,
            "path": path,
            "count": len(detailed_items),
            "items": detailed_items,
            "message": f"Listed {len(detailed_items)} items in {path}"
        }

    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": f"Error listing directory: {str(e)}"}


def _write_file_logic(
    path: str,
    content: str,
    mode: str = "overwrite",
    max_bytes: int = 200_000
) -> Dict[str, Any]:
    """Writes text content to a file."""
    try:
        path = os.path.abspath(os.path.expanduser(path))

        if not isinstance(content, str):
            return {"error": "content must be a string"}
        if len(content.encode("utf-8")) > max_bytes:
            return {
                "error": f"Content too large ({len(content)} bytes). "
                         f"Max allowed is {max_bytes} bytes."
            }

        mode = mode.lower()
        if mode not in ["overwrite", "append"]:
            return {"error": "mode must be 'overwrite' or 'append'"}

        parent_dir = os.path.dirname(path)
        if parent_dir and not os.path.exists(parent_dir):
            if ".." in parent_dir:
                return {"error": "Directory traversal not allowed"}
            os.makedirs(parent_dir, exist_ok=True)

        file_mode = "w" if mode == "overwrite" else "a"

        with open(path, file_mode, encoding="utf-8") as f:
            written = f.write(content)

        return {
            "success": True,
            "path": path,
            "mode": mode,
            "bytes_written": written,
            "message": f"Wrote {written} bytes to {path} ({mode})"
        }

    except PermissionError:
        return {"error": f"Permission denied: {path}"}
    except Exception as e:
        return {"error": f"Error writing file: {str(e)}"}


def register(server):
    @server.tool()
    def read_file_meta(
        # --- Mandatory read_file parameters ---
        source_path: str,
        source_start_line: Optional[int] = None,
        source_end_line: Optional[int] = None,
        source_max_bytes: int = 200_000,

        # --- Optional add_paragraph parameters ---
        add_paragraph_doc_path: Optional[str] = None,
        add_paragraph_text: Optional[str] = None,
        add_paragraph_style: Optional[str] = None,

        # --- Optional execute_command parameters ---
        command_to_execute: Optional[str] = None,
        command_timeout: int = 10,

        # --- Optional list_directory parameters ---
        directory_to_list: Optional[str] = None,

        # --- Optional write_file parameters ---
        write_output_path: Optional[str] = None,
        write_content: Optional[str] = None,
        write_mode: str = "overwrite",
        write_max_bytes: int = 200_000,

        # --- Optional second read_file parameters ---
        second_read_path: Optional[str] = None,
        second_read_start_line: Optional[int] = None,
        second_read_end_line: Optional[int] = None,
        second_read_max_bytes: int = 200_000,

    ) -> Dict[str, Any]:
        """
        Reads a file and optionally performs additional file system or document operations.

        Mandatory behavior (FIRST TOOL: read_file):
            - Always reads the content of `source_path`.

        Optional behaviors, controlled by parameters:
            - `add_paragraph`: If `add_paragraph_doc_path` and `add_paragraph_text` are provided, adds a paragraph to a Word document.
            - `execute_command`: If `command_to_execute` is provided, executes a whitelisted shell command.
            - `list_directory`: If `directory_to_list` is provided, lists its contents.
            - `write_file`: If `write_output_path` and `write_content` are provided, writes content to a file.
            - `read_file` (second instance): If `second_read_path` is provided, performs another file read operation.

        Args:
            source_path: Path to the file to read (mandatory first operation).
            source_start_line: Optional starting line (1-based) for the primary read.
            source_end_line: Optional ending line (1-based) for the primary read.
            source_max_bytes: Safety limit for primary file size.

            add_paragraph_doc_path: Path to a .docx file for adding a paragraph. If provided, `add_paragraph_text` is required.
            add_paragraph_text: Text content for the paragraph. Required if `add_paragraph_doc_path` is set.
            add_paragraph_style: Optional paragraph style name for the Word document.

            command_to_execute: A shell command string to execute (e.g., "ls -l").
            command_timeout: Maximum execution time for the command in seconds.

            directory_to_list: Path to a directory to list its contents.

            write_output_path: Path to the file where content will be written. If provided, `write_content` is required.
            write_content: Text content to write to the file. Required if `write_output_path` is set.
            write_mode: 'overwrite' or 'append' for writing the file.
            write_max_bytes: Max allowed content size for writing.

            second_read_path: Path to a file for a second, optional read operation.
            second_read_start_line: Optional starting line (1-based) for the second read.
            second_read_end_line: Optional ending line (1-based) for the second read.
            second_read_max_bytes: Safety limit for second file size.

        Returns:
            A dict describing all operations performed, their success status, and any errors.
        """
        results: Dict[str, Any] = {
            "success": True,
            "overall_message": "All requested operations completed successfully.",
            "operations": {},
        }
        error_messages: List[str] = []

        # --- Mandatory Operation: primary read_file ---
        primary_read_file_result = _read_file_logic(
            source_path, source_start_line, source_end_line, source_max_bytes
        )
        results["operations"]["primary_read_file"] = primary_read_file_result
        if "error" in primary_read_file_result:
            results["success"] = False
            results["overall_message"] = f"Primary file read failed: {primary_read_file_result['error']}"
            return results # Mandatory step failed, return immediately

        # --- Optional Operation: add_paragraph ---
        if add_paragraph_doc_path:
            if add_paragraph_text is None:
                ap_error = {"error": "add_paragraph_text is required if add_paragraph_doc_path is provided."}
                results["operations"]["add_paragraph"] = ap_error
                error_messages.append(ap_error["error"])
            else:
                add_paragraph_result = _add_paragraph_logic(
                    add_paragraph_doc_path, add_paragraph_text, add_paragraph_style
                )
                results["operations"]["add_paragraph"] = add_paragraph_result
                if "error" in add_paragraph_result:
                    error_messages.append(f"Add paragraph failed: {add_paragraph_result['error']}")

        # --- Optional Operation: execute_command ---
        if command_to_execute:
            execute_command_result = _execute_command_logic(
                command_to_execute, command_timeout
            )
            results["operations"]["execute_command"] = execute_command_result
            if "error" in execute_command_result:
                error_messages.append(f"Command execution failed: {execute_command_result['error']}")

        # --- Optional Operation: list_directory ---
        if directory_to_list:
            list_directory_result = _list_directory_logic(directory_to_list)
            results["operations"]["list_directory"] = list_directory_result
            if "error" in list_directory_result:
                error_messages.append(f"Directory listing failed: {list_directory_result['error']}")

        # --- Optional Operation: write_file ---
        if write_output_path:
            if write_content is None:
                wf_error = {"error": "write_content is required if write_output_path is provided."}
                results["operations"]["write_file"] = wf_error
                error_messages.append(wf_error["error"])
            else:
                write_file_result = _write_file_logic(
                    write_output_path, write_content, write_mode, write_max_bytes
                )
                results["operations"]["write_file"] = write_file_result
                if "error" in write_file_result:
                    error_messages.append(f"Write file failed: {write_file_result['error']}")

        # --- Optional Operation: second read_file ---
        if second_read_path:
            second_read_file_result = _read_file_logic(
                second_read_path, second_read_start_line, second_read_end_line, second_read_max_bytes
            )
            results["operations"]["second_read_file"] = second_read_file_result
            if "error" in second_read_file_result:
                error_messages.append(f"Second file read failed: {second_read_file_result['error']}")

        if error_messages:
            results["success"] = False
            results["overall_message"] = "Some operations failed: " + "; ".join(error_messages)

        return results
```