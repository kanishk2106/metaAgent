from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os
import re
from typing import Optional, Dict, Any, List
def register(server):
    @server.tool()
    def add_heading_and_document_elements(
        filename: str,
        heading_text: str,
        heading_level: int = 1,
        add_paragraph_text: Optional[str] = None,
        add_paragraph_style: Optional[str] = None,
        add_table_rows: Optional[int] = None,
        add_table_cols: Optional[int] = None,
        add_table_data: Optional[List[List[Any]]] = None,
        format_target_table_index: Optional[int] = None,
        format_has_header_row: bool = False,
        format_border_style: Optional[str] = None,
        format_shading: Optional[List[Dict[str, Any]]] = None,
    ) -> Dict[str, Any]:
        try:
            filename = os.path.abspath(filename)
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith(".docx"):
                return {"error": "File must be a .docx file"}
            if not isinstance(heading_level, int) or heading_level < 1 or heading_level > 9:
                return {"error": "heading_level must be an integer between 1 and 9"}
            add_table_enabled = add_table_rows is not None and add_table_cols is not None
            if add_table_enabled:
                if not isinstance(add_table_rows, int) or not isinstance(add_table_cols, int):
                    return {"error": "add_table_rows and add_table_cols must be integers"}
                if add_table_rows < 1 or add_table_cols < 1:
                    return {"error": "add_table_rows and add_table_cols must be >= 1"}
                if add_table_data is not None:
                    if not isinstance(add_table_data, list):
                        return {"error": "add_table_data must be a 2D list of lists"}
                    for r in add_table_data:
                        if not isinstance(r, list):
                            return {"error": "Each row in add_table_data must be a list"}
            doc = Document(filename)
            initial_table_count = len(doc.tables)
            operations = {
                "heading_added": False,
                "paragraph_added": False,
                "table_added": False,
                "table_formatted": False,
                "formatted_table_index": None,
            }
            last_added_table_index: Optional[int] = None
            doc.add_heading(heading_text, level=heading_level)
            operations["heading_added"] = True
            if add_paragraph_text is not None:
                paragraph = doc.add_paragraph(add_paragraph_text)
                if add_paragraph_style:
                    try:
                        paragraph.style = add_paragraph_style
                    except KeyError:
                        return {
                            "error": f"Paragraph style '{add_paragraph_style}' not found in document",
                            "filename": filename,
                        }
                operations["paragraph_added"] = True
            if add_table_enabled:
                table = doc.add_table(rows=add_table_rows, cols=add_table_cols)
                populated_rows = 0
                if add_table_data:
                    for i, row_data in enumerate(add_table_data):
                        if i >= add_table_rows:
                            break
                        for j, cell_value in enumerate(row_data):
                            if j >= add_table_cols:
                                break
                            table.rows[i].cells[j].text = str(cell_value)
                        populated_rows += 1
                operations["table_added"] = True
                operations["table_rows"] = add_table_rows
                operations["table_cols"] = add_table_cols
                operations["table_populated_rows"] = populated_rows
                last_added_table_index = initial_table_count # New table is at this index
            target_idx_for_formatting = None
            if format_target_table_index is not None:
                target_idx_for_formatting = format_target_table_index
            elif operations["table_added"] and (format_has_header_row or format_border_style or format_shading):
                target_idx_for_formatting = last_added_table_index
            if target_idx_for_formatting is not None:
                if target_idx_for_formatting < 0 or target_idx_for_formatting >= len(doc.tables):
                    return {
                        "error": f"Table index {target_idx_for_formatting} out of range for formatting "
                                 f"(document has {len(doc.tables)} tables, indices 0-{len(doc.tables)-1})",
                        "filename": filename,
                    }
                table_to_format = doc.tables[target_idx_for_formatting]
                if format_has_header_row and len(table_to_format.rows) > 0:
                    try:
                        table_to_format.style = "Table Grid" # Common default style for header rows
                    except Exception:
                        pass # Ignore if style not found, not critical for basic formatting
                allowed_borders = ["single", "double", "none", None]
                if format_border_style not in allowed_borders:
                    return {
                        "error": f"Unsupported format_border_style '{format_border_style}'. Allowed: {allowed_borders}",
                        "filename": filename,
                        "formatted_table_index": target_idx_for_formatting,
                    }
                shaded_rows_count = 0
                if format_shading:
                    if not isinstance(format_shading, list):
                        return {
                            "error": "format_shading must be a list of dict entries",
                            "filename": filename,
                            "formatted_table_index": target_idx_for_formatting,
                        }
                    hex_pattern = re.compile(r"^[0-9A-Fa-f]{6}$")
                    for entry in format_shading:
                        if not isinstance(entry, dict):
                            continue # Skip malformed entries
                        row_idx = entry.get("row")
                        color = entry.get("color", "FFFFFF") # Default to white if no color specified
                        if not isinstance(row_idx, int) or row_idx < 0 or row_idx >= len(table_to_format.rows):
                            continue # Skip invalid row indices
                        if not isinstance(color, str) or not hex_pattern.match(color):
                            continue # Skip invalid hex colors
                        try:
                            for cell in table_to_format.rows[row_idx].cells:
                                shade = OxmlElement("w:shd")
                                shade.set(qn("w:fill"), color)
                                cell._element.get_or_add_tcPr().append(shade)
                            shaded_rows_count += 1
                        except Exception:
                            continue # Silently fail on individual cell shading error
                operations["table_formatted"] = True
                operations["formatted_table_index"] = target_idx_for_formatting
                operations["format_has_header_row"] = format_has_header_row
                operations["format_border_style"] = format_border_style
                operations["format_shaded_rows"] = shaded_rows_count
            doc.save(filename)
            message_parts = [f"Added heading '{heading_text}' (level {heading_level})."]
            if operations["paragraph_added"]:
                message_parts.append("Added paragraph.")
            if operations["table_added"]:
                message_parts.append(f"Added {operations['table_rows']}x{operations['table_cols']} table.")
            if operations["table_formatted"]:
                message_parts.append(f"Formatted table at index {operations['formatted_table_index']}.")
            return {
                "success": True,
                "filename": filename,
                "operations_summary": operations,
                "message": " ".join(message_parts),
            }
        except Exception as e:
            return {"error": f"Failed to perform document operations: {str(e)}"}