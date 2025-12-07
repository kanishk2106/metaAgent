from docx import Document
import os
from typing import Optional, Dict, Any, List
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import re
def register(server):
    @server.tool()
    def add_paragraph_meta(
        filename: str,
        paragraph_text: str,
        paragraph_style: Optional[str] = None,
        add_heading_text: Optional[str] = None,  # If provided, a heading is added
        add_heading_level: int = 1,              # Level for the optional heading (1-9)
        add_additional_paragraph_text: Optional[str] = None,  # If provided, another paragraph is added
        add_additional_paragraph_style: Optional[str] = None, # Optional style for the second paragraph
        add_table_rows: Optional[int] = None,  # If provided (and >0), a table is added
        add_table_cols: Optional[int] = None,  # Requires add_table_rows
        add_table_data: Optional[List[List[Any]]] = None, # Optional 2D list of cell data for the table
        apply_table_formatting: bool = False,  # Flag to enable any table formatting
        format_table_index: Optional[int] = None, # Index of the table to format (0-based).
        format_has_header_row: bool = False,
        format_border_style: Optional[str] = None, # E.g., 'single', 'double', 'none'
        format_shading: Optional[List[Dict[str, Any]]] = None, # List of {"row": int, "color": "HEX"}
    ) -> Dict[str, Any]:
        operations_summary = {
            "heading_added": False,
            "mandatory_paragraph_added": False,
            "additional_paragraph_added": False,
            "table_added": False,
            "table_formatted": False,
            "table_formatted_index": None,
        }
        try:
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith(".docx"):
                return {"error": "File must be a .docx file"}
            if add_heading_text is not None:
                if not isinstance(add_heading_level, int) or add_heading_level < 1 or add_heading_level > 9:
                    return {"error": "add_heading_level must be an integer between 1 and 9."}
            if add_table_rows is not None:
                if not isinstance(add_table_rows, int) or add_table_rows < 1:
                    return {"error": "add_table_rows must be an integer >= 1 if provided."}
                if add_table_cols is None:
                    return {"error": "add_table_cols is required if add_table_rows is provided."}
                if not isinstance(add_table_cols, int) or add_table_cols < 1:
                    return {"error": "add_table_cols must be an integer >= 1 if provided."}
                if add_table_data is not None:
                    if not isinstance(add_table_data, list):
                        return {"error": "add_table_data must be a 2D list of lists."}
                    for r_data in add_table_data:
                        if not isinstance(r_data, list):
                            return {"error": "Each row in add_table_data must be a list."}
            if apply_table_formatting:
                allowed_borders = ["single", "double", "none", None]
                if format_border_style not in allowed_borders:
                    return {
                        "error": f"Unsupported format_border_style '{format_border_style}'. Allowed: {allowed_borders}"
                    }
                if format_shading is not None and not isinstance(format_shading, list):
                    return {"error": "format_shading must be a list of dict entries."}
            doc = Document(filename)
            last_added_table_index = None
            if add_heading_text is not None:
                doc.add_heading(add_heading_text, level=add_heading_level)
                operations_summary["heading_added"] = True
            paragraph_elem = doc.add_paragraph(paragraph_text)
            if paragraph_style:
                try:
                    paragraph_elem.style = paragraph_style
                except KeyError:
                    return {
                        "error": f"Style '{paragraph_style}' not found for mandatory paragraph.",
                        "filename": filename,
                        "operations": operations_summary
                    }
            operations_summary["mandatory_paragraph_added"] = True
            if add_additional_paragraph_text is not None:
                additional_paragraph_elem = doc.add_paragraph(add_additional_paragraph_text)
                if add_additional_paragraph_style:
                    try:
                        additional_paragraph_elem.style = add_additional_paragraph_style
                    except KeyError:
                        return {
                            "error": f"Style '{add_additional_paragraph_style}' not found for additional paragraph.",
                            "filename": filename,
                            "operations": operations_summary
                        }
                operations_summary["additional_paragraph_added"] = True
            if add_table_rows is not None and add_table_cols is not None:
                table = doc.add_table(rows=add_table_rows, cols=add_table_cols)
                populated_rows_count = 0
                if add_table_data:
                    for i, row_data in enumerate(add_table_data):
                        if i >= add_table_rows:
                            break
                        for j, cell_value in enumerate(row_data):
                            if j >= add_table_cols:
                                break
                            table.rows[i].cells[j].text = str(cell_value)
                        populated_rows_count += 1
                operations_summary["table_added"] = True
                last_added_table_index = len(doc.tables) - 1 # Store index of this newly added table
                operations_summary["table_rows"] = add_table_rows
                operations_summary["table_cols"] = add_table_cols
                operations_summary["table_populated_rows"] = populated_rows_count
            if apply_table_formatting:
                target_table_index = format_table_index
                if target_table_index is None and operations_summary["table_added"]:
                    target_table_index = last_added_table_index
                if target_table_index is not None:
                    if target_table_index < 0 or target_table_index >= len(doc.tables):
                        return {
                            "error": f"Table index {target_table_index} out of range (0-{len(doc.tables)-1}) for formatting.",
                            "filename": filename,
                            "operations": operations_summary
                        }
                    table_to_format = doc.tables[target_table_index]
                    if format_has_header_row and len(table_to_format.rows) > 0:
                        try:
                            table_to_format.style = "Table Grid"
                        except Exception:
                            pass # Style might not exist, don't crash the tool
                    shaded_rows_applied_count = 0
                    if format_shading:
                        hex_pattern = re.compile(r"^[0-9A-Fa-f]{6}$")
                        for entry in format_shading:
                            if not isinstance(entry, dict):
                                continue # Skip invalid entries
                            row_idx = entry.get("row")
                            color = entry.get("color", "FFFFFF")
                            if not isinstance(row_idx, int) or row_idx < 0 or row_idx >= len(table_to_format.rows):
                                continue # Skip invalid row index
                            if not isinstance(color, str) or not hex_pattern.match(color):
                                continue # Skip invalid color hex code
                            try:
                                for cell in table_to_format.rows[row_idx].cells:
                                    shade = OxmlElement("w:shd")
                                    shade.set(qn("w:fill"), color)
                                    cell._element.get_or_add_tcPr().append(shade)
                                shaded_rows_applied_count += 1
                            except Exception:
                                pass # Continue if shading fails for a specific cell/row
                    operations_summary["table_formatted"] = True
                    operations_summary["table_formatted_index"] = target_table_index
                    operations_summary["table_formatted_has_header_row"] = format_has_header_row
                    operations_summary["table_formatted_border_style"] = format_border_style
                    operations_summary["table_formatted_shaded_rows"] = shaded_rows_applied_count
                elif apply_table_formatting: # apply_table_formatting is True, but no target_table_index was resolved
                     return {
                        "error": "Table formatting requested, but no table was newly added and no specific 'format_table_index' was provided to format an existing table.",
                        "filename": filename,
                        "operations": operations_summary
                     }
            doc.save(filename)
            final_message_parts = ["Document updated:"]
            if operations_summary["mandatory_paragraph_added"]:
                final_message_parts.append("added primary paragraph")
            if operations_summary["heading_added"]:
                final_message_parts.append("added heading")
            if operations_summary["additional_paragraph_added"]:
                final_message_parts.append("added additional paragraph")
            if operations_summary["table_added"]:
                final_message_parts.append(f"added {operations_summary['table_rows']}x{operations_summary['table_cols']} table")
            if operations_summary["table_formatted"]:
                final_message_parts.append(f"formatted table at index {operations_summary['table_formatted_index']}")
            final_message = " ".join(final_message_parts) + "."
            return {
                "success": True,
                "filename": filename,
                "message": final_message,
                "operations": operations_summary,
            }
        except Exception as e:
            return {
                "error": f"Failed to perform document operations: {str(e)}",
                "filename": filename,
                "operations": operations_summary
            }