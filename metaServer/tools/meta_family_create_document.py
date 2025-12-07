from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
from typing import Optional, Dict, Any, List
def register(server):
    @server.tool()
    def create_document_with_content(
        filename: str,
        title: Optional[str] = None,
        author: Optional[str] = None,
        heading_text: Optional[str] = None,
        heading_level: int = 1,
        paragraph_text: Optional[str] = None,
        paragraph_style: Optional[str] = None,
        table_rows: Optional[int] = None,
        table_cols: Optional[int] = None,
        table_data: Optional[List[List[Any]]] = None,
        format_table_index: Optional[int] = None,
        format_table_has_header_row: bool = False,
        format_table_border_style: Optional[str] = None,
        format_table_shading: Optional[List[Dict[str, Any]]] = None,
    ) -> Dict[str, Any]:
        operations = {
            "document_created": False,
            "filename": os.path.abspath(filename),
            "title": title if title else "",
            "author": author if author else "",
            "heading_added": False,
            "paragraph_added": False,
            "table_added": False,
            "table_index_added": None, # Will store the index of the table if added by this tool
            "table_formatted": False,
            "table_formatted_index": None, # Will store the index of the table that was formatted
        }
        try:
            filename = os.path.abspath(filename)
            if not filename.endswith(".docx"):
                return {"error": "Filename must end with .docx"}
            if os.path.exists(filename):
                return {"error": f"File already exists: {filename}"}
            parent_dir = os.path.dirname(filename)
            if parent_dir and parent_dir.strip() and not os.path.exists(parent_dir):
                os.makedirs(parent_dir, exist_ok=True)
            doc = Document()
            core = doc.core_properties
            if title:
                core.title = title
            if author:
                core.author = author
            operations["document_created"] = True
            if heading_text is not None:
                if not isinstance(heading_level, int) or heading_level < 1 or heading_level > 9:
                    return {"error": "heading_level must be an integer between 1 and 9"}
                doc.add_heading(heading_text, level=heading_level)
                operations["heading_added"] = True
                operations["heading_text"] = heading_text
                operations["heading_level"] = heading_level
            if paragraph_text is not None:
                paragraph = doc.add_paragraph(paragraph_text)
                if paragraph_style:
                    try:
                        paragraph.style = paragraph_style
                    except KeyError:
                        return {
                            "error": f"Paragraph style '{paragraph_style}' not found in document",
                            "filename": filename,
                        }
                operations["paragraph_added"] = True
                operations["paragraph_text_length"] = len(paragraph_text)
                operations["paragraph_style"] = paragraph_style or "default"
            if table_rows is not None and table_cols is not None:
                if not isinstance(table_rows, int) or not isinstance(table_cols, int):
                    return {"error": "table_rows and table_cols must be integers"}
                if table_rows < 1 or table_cols < 1:
                    return {"error": "table_rows and table_cols must be >= 1"}
                if table_data is not None:
                    if not isinstance(table_data, list):
                        return {"error": "table_data must be a 2D list of lists"}
                    for r in table_data:
                        if not isinstance(r, list):
                            return {"error": "Each row in table_data must be a list"}
                table = doc.add_table(rows=table_rows, cols=table_cols)
                operations["table_index_added"] = len(doc.tables) - 1
                populated_rows = 0
                if table_data:
                    for i, row_data in enumerate(table_data):
                        if i >= table_rows:
                            break
                        for j, cell_value in enumerate(row_data):
                            if j >= table_cols:
                                break
                            table.rows[i].cells[j].text = str(cell_value)
                        populated_rows += 1
                operations["table_added"] = True
                operations["table_rows"] = table_rows
                operations["table_cols"] = table_cols
                operations["table_populated_rows"] = populated_rows
            actual_format_table_index = None
            if format_table_index is not None:
                actual_format_table_index = format_table_index
            elif operations["table_added"]: # If a table was just added, and no specific index for format_table was given, use that one.
                actual_format_table_index = operations["table_index_added"]
            if actual_format_table_index is not None and (
                format_table_has_header_row
                or format_table_border_style
                or format_table_shading
            ):
                if actual_format_table_index < 0 or actual_format_table_index >= len(doc.tables):
                    return {
                        "error": f"Table index {actual_format_table_index} for formatting out of range (0-{len(doc.tables)-1})",
                        "filename": filename,
                    }
                table_to_format = doc.tables[actual_format_table_index]
                operations["table_formatted_index"] = actual_format_table_index
                if format_table_has_header_row and len(table_to_format.rows) > 0:
                    try:
                        table_to_format.style = "Table Grid"
                    except Exception:
                        pass # Ignore if style not found, don't crash the meta tool.
                    operations["format_table_has_header_row"] = True
                allowed_borders = ["single", "double", "none", None]
                if format_table_border_style is not None and format_table_border_style not in allowed_borders:
                    return {
                        "error": f"Unsupported format_table_border_style '{format_table_border_style}'. Allowed: {allowed_borders}"
                    }
                if format_table_border_style:
                    operations["format_table_border_style"] = format_table_border_style
                shaded_rows_count = 0
                if format_table_shading:
                    if not isinstance(format_table_shading, list):
                        return {"error": "format_table_shading must be a list of dict entries"}
                    hex_pattern = re.compile(r"^[0-9A-Fa-f]{6}$")
                    for entry in format_table_shading:
                        if not isinstance(entry, dict):
                            continue # Skip malformed entries
                        row_idx = entry.get("row")
                        color = entry.get("color", "FFFFFF") # Default to white if not specified
                        if not isinstance(row_idx, int) or row_idx < 0 or row_idx >= len(table_to_format.rows):
                            continue # Skip invalid row indices
                        if not isinstance(color, str) or not hex_pattern.match(color):
                            continue # Skip invalid hex colors
                        try:
                            for cell in table_to_format.rows[row_idx].cells:
                                shade = OxmlElement("w:shd")
                                shade.set(qn("w:fill"), color)
                                cell._element.get_or_add_tcPr().append(shade)
                            shaded_rows_count += 1
                        except Exception:
                            continue # Ignore individual shading errors
                if shaded_rows_count > 0:
                    operations["format_table_shading_rows"] = shaded_rows_count
                operations["table_formatted"] = True
            doc.save(filename)
            message_parts = [f"Created new Word document at {filename}."]
            if operations["heading_added"]:
                message_parts.append(f"Added heading '{heading_text}'.")
            if operations["paragraph_added"]:
                message_parts.append("Added paragraph.")
            if operations["table_added"]:
                message_parts.append(f"Added {operations['table_rows']}x{operations['table_cols']} table.")
            if operations["table_formatted"]:
                message_parts.append(f"Formatted table at index {operations['table_formatted_index']}.")
            return {
                "success": True,
                "filename": filename,
                "operations": operations,
                "message": " ".join(message_parts)
            }
        except Exception as e:
            return {"error": f"Failed to create document with content: {str(e)}"}