from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os
import re
from typing import Optional, Dict, Any, List
def register(server):
    @server.tool()
    def add_table_meta(
        filename: str,
        table_rows: int,
        table_cols: int,
        table_data: Optional[List[List[Any]]] = None,
        add_heading_text: Optional[str] = None,
        heading_level: int = 1,
        add_paragraph_text: Optional[str] = None,
        paragraph_style: Optional[str] = None,
        add_another_table_rows: Optional[int] = None,
        another_table_cols: Optional[int] = None,
        another_table_data: Optional[List[List[Any]]] = None,
        format_table_index: Optional[int] = None,
        format_has_header_row: bool = False,
        format_border_style: Optional[str] = None,
        format_shading: Optional[List[Dict[str, Any]]] = None,
    ) -> Dict[str, Any]:
        try:
            filename = os.path.abspath(filename)
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith(".docx"):
                return {"error": "File must be a .docx file"}
            if not isinstance(table_rows, int) or not isinstance(table_cols, int):
                return {"error": "Primary table rows and columns must be integers"}
            if table_rows < 1 or table_cols < 1:
                return {"error": "Primary table rows and columns must be >= 1"}
            if table_data is not None:
                if not isinstance(table_data, list):
                    return {"error": "Primary table data must be a 2D list of lists"}
                for r in table_data:
                    if not isinstance(r, list):
                        return {"error": "Each row in primary table data must be a list"}
            if add_heading_text is not None:
                if not isinstance(heading_level, int) or heading_level < 1 or heading_level > 9:
                    return {"error": "heading_level must be an integer between 1 and 9"}
            if add_another_table_rows is not None:
                if another_table_cols is None:
                    return {"error": "another_table_cols is mandatory if add_another_table_rows is provided"}
                if not isinstance(add_another_table_rows, int) or not isinstance(another_table_cols, int):
                    return {"error": "Second table rows and columns must be integers"}
                if add_another_table_rows < 1 or another_table_cols < 1:
                    return {"error": "Second table rows and columns must be >= 1"}
                if another_table_data is not None:
                    if not isinstance(another_table_data, list):
                        return {"error": "Second table data must be a 2D list of lists"}
                    for r in another_table_data:
                        if not isinstance(r, list):
                            return {"error": "Each row in second table data must be a list"}
            doc = Document(filename)
            initial_table_count = len(doc.tables)
            operations = {
                "filename": filename,
                "primary_table_added": False,
                "heading_added": False,
                "paragraph_added": False,
                "second_table_added": False,
                "table_formatted": False,
                "formatted_table_index": None,
                "message": [],
            }
            primary_table = doc.add_table(rows=table_rows, cols=table_cols)
            populated_primary_rows = 0
            if table_data:
                for i, row_data in enumerate(table_data):
                    if i >= table_rows: break
                    for j, cell_value in enumerate(row_data):
                        if j >= table_cols: break
                        primary_table.rows[i].cells[j].text = str(cell_value)
                    populated_primary_rows += 1
            operations["primary_table_added"] = True
            operations["primary_table_details"] = {
                "rows": table_rows, "cols": table_cols, "populated_rows": populated_primary_rows
            }
            operations["message"].append(f"Added primary {table_rows}x{table_cols} table.")
            if add_heading_text is not None:
                doc.add_heading(add_heading_text, level=heading_level)
                operations["heading_added"] = True
                operations["heading_details"] = {"text": add_heading_text, "level": heading_level}
                operations["message"].append(f"Added heading '{add_heading_text}' (level {heading_level}).")
            if add_paragraph_text is not None:
                paragraph = doc.add_paragraph(add_paragraph_text)
                if paragraph_style:
                    try:
                        paragraph.style = paragraph_style
                    except KeyError:
                        return {
                            "error": f"Style '{paragraph_style}' not found in document for paragraph",
                            "filename": filename,
                        }
                operations["paragraph_added"] = True
                operations["paragraph_details"] = {"text_length": len(add_paragraph_text), "style": paragraph_style or "default"}
                operations["message"].append(f"Added paragraph with {len(add_paragraph_text)} characters.")
            if add_another_table_rows is not None:
                another_table = doc.add_table(rows=add_another_table_rows, cols=another_table_cols)
                populated_another_rows = 0
                if another_table_data:
                    for i, row_data in enumerate(another_table_data):
                        if i >= add_another_table_rows: break
                        for j, cell_value in enumerate(row_data):
                            if j >= another_table_cols: break
                            another_table.rows[i].cells[j].text = str(cell_value)
                        populated_another_rows += 1
                operations["second_table_added"] = True
                operations["second_table_details"] = {
                    "rows": add_another_table_rows, "cols": another_table_cols, "populated_rows": populated_another_rows
                }
                operations["message"].append(f"Added second {add_another_table_rows}x{another_table_cols} table.")
            if format_table_index is not None:
                if format_table_index < 0 or format_table_index >= len(doc.tables):
                    return {
                        "error": f"Table index {format_table_index} out of range (0-{len(doc.tables)-1}) for formatting",
                        "filename": filename,
                    }
                table_to_format = doc.tables[format_table_index]
                if format_has_header_row and len(table_to_format.rows) > 0:
                    try:
                        table_to_format.style = "Table Grid" # Common style, may fail if not present
                    except Exception:
                        pass # Silently fail if style not found, as per base tool
                allowed_borders = ["single", "double", "none", None]
                if format_border_style not in allowed_borders:
                    return {
                        "error": f"Unsupported border_style '{format_border_style}'. Allowed: {allowed_borders}",
                        "filename": filename,
                    }
                shaded_rows_count = 0
                if format_shading:
                    if not isinstance(format_shading, list):
                        return {"error": "format_shading must be a list of dict entries", "filename": filename}
                    hex_pattern = re.compile(r"^[0-9A-Fa-f]{6}$")
                    for entry in format_shading:
                        if not isinstance(entry, dict):
                            continue # Skip malformed entries
                        row_idx = entry.get("row")
                        color = entry.get("color", "FFFFFF")
                        if not isinstance(row_idx, int) or row_idx < 0 or row_idx >= len(table_to_format.rows):
                            continue # Skip invalid row index
                        if not isinstance(color, str) or not hex_pattern.match(color):
                            continue # Skip invalid color
                        try:
                            for cell in table_to_format.rows[row_idx].cells:
                                shade = OxmlElement("w:shd")
                                shade.set(qn("w:fill"), color)
                                cell._element.get_or_add_tcPr().append(shade)
                            shaded_rows_count += 1
                        except Exception:
                            continue # Silently fail for individual shading errors
                operations["table_formatted"] = True
                operations["formatted_table_index"] = format_table_index
                operations["formatting_details"] = {
                    "has_header_row": format_has_header_row,
                    "border_style": format_border_style,
                    "shading_rows_applied": shaded_rows_count,
                }
                operations["message"].append(f"Formatted table at index {format_table_index} (shaded {shaded_rows_count} rows).")
            doc.save(filename)
            return {
                "success": True,
                "overall_message": " | ".join(operations["message"]) or "No operations performed beyond mandatory table addition.",
                "operations_summary": {
                    "filename": operations["filename"],
                    "primary_table_added": operations["primary_table_added"],
                    "primary_table_details": operations.get("primary_table_details"),
                    "heading_added": operations["heading_added"],
                    "heading_details": operations.get("heading_details"),
                    "paragraph_added": operations["paragraph_added"],
                    "paragraph_details": operations.get("paragraph_details"),
                    "second_table_added": operations["second_table_added"],
                    "second_table_details": operations.get("second_table_details"),
                    "table_formatted": operations["table_formatted"],
                    "formatted_table_index": operations["formatted_table_index"],
                    "formatting_details": operations.get("formatting_details"),
                },
            }
        except Exception as e:
            return {"error": f"Failed during add_table_meta operation: {str(e)}"}