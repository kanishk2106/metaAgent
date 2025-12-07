from docx import Document
import os
def register(server):
    @server.tool()
    def add_paragraph(filename: str, text: str, style: str = None) -> dict:
        try:
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith('.docx'):
                return {"error": "File must be a .docx file"}
            doc = Document(filename)
            paragraph = doc.add_paragraph(text)
            if style:
                try:
                    paragraph.style = style
                except KeyError:
                    return {"error": f"Style '{style}' not found in document"}
            doc.save(filename)
            return {
                "success": True,
                "filename": filename,
                "text": text,
                "style": style if style else "default",
                "message": f"Added paragraph with {len(text)} characters"
            }
        except Exception as e:
            return {"error": f"Failed to add paragraph: {str(e)}"}
