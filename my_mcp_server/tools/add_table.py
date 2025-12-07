from docx import Document
import os
def register(server):
    @server.tool()
    def add_table(
        filename: str,
        rows: int,
        cols: int,
        data: list | None = None
    ) -> dict:
        try:
            filename = os.path.abspath(filename)
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith(".docx"):
                return {"error": "File must be a .docx file"}
            if not isinstance(rows, int) or not isinstance(cols, int):
                return {"error": "Rows and columns must be integers"}
            if rows < 1 or cols < 1:
                return {"error": "Rows and columns must be >= 1"}
            if data is not None:
                if not isinstance(data, list):
                    return {"error": "data must be a 2D list of lists"}
                for r in data:
                    if not isinstance(r, list):
                        return {"error": "Each row in data must be a list"}
            doc = Document(filename)
            table = doc.add_table(rows=rows, cols=cols)
            populated_rows = 0
            if data:
                for i, row_data in enumerate(data):
                    if i >= rows:
                        break
                    for j, cell_value in enumerate(row_data):
                        if j >= cols:
                            break
                        table.rows[i].cells[j].text = str(cell_value)
                    populated_rows += 1
            doc.save(filename)
            return {
                "success": True,
                "filename": filename,
                "rows": rows,
                "cols": cols,
                "populated_rows": populated_rows,
                "message": f"Added {rows}x{cols} table with {populated_rows} populated rows"
            }
        except Exception as e:
            return {"error": f"Failed to add table: {str(e)}"}
