from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os
import re
def register(server):
    @server.tool()
    def format_table(
        filename: str,
        table_index: int,
        has_header_row: bool = False,
        border_style: str = None,
        shading: list | None = None
    ) -> dict:
        try:
            filename = os.path.abspath(filename)
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith(".docx"):
                return {"error": "File must be a .docx file"}
            doc = Document(filename)
            if table_index < 0 or table_index >= len(doc.tables):
                return {
                    "error": f"Table index {table_index} out of range (0-{len(doc.tables)-1})"
                }
            table = doc.tables[table_index]
            if has_header_row and len(table.rows) > 0:
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass  # incorrect style name should not crash the tool
            allowed_borders = ["single", "double", "none", None]
            if border_style not in allowed_borders:
                return {
                    "error": f"Unsupported border_style '{border_style}'. Allowed: {allowed_borders}"
                }
            shaded_rows = 0
            if shading:
                if not isinstance(shading, list):
                    return {"error": "shading must be a list of dict entries"}
                hex_pattern = re.compile(r"^[0-9A-Fa-f]{6}$")
                for entry in shading:
                    if not isinstance(entry, dict):
                        continue
                    row_idx = entry.get("row")
                    color = entry.get("color", "FFFFFF")
                    if not isinstance(row_idx, int) or row_idx < 0 or row_idx >= len(table.rows):
                        continue
                    if not isinstance(color, str) or not hex_pattern.match(color):
                        continue
                    try:
                        for cell in table.rows[row_idx].cells:
                            shade = OxmlElement("w:shd")
                            shade.set(qn("w:fill"), color)
                            cell._element.get_or_add_tcPr().append(shade)
                        shaded_rows += 1
                    except Exception:
                        continue
            doc.save(filename)
            return {
                "success": True,
                "filename": filename,
                "table_index": table_index,
                "has_header_row": has_header_row,
                "shading_rows": shaded_rows,
                "border_style": border_style,
                "message": f"Formatted table {table_index} (shaded {shaded_rows} rows)"
            }
        except Exception as e:
            return {"error": f"Failed to format table: {str(e)}"}
