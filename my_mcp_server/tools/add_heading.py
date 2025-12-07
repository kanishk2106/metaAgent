from docx import Document
import os
def register(server):
    @server.tool()
    def add_heading(filename: str, text: str, level: int = 1) -> dict:
        try:
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith('.docx'):
                return {"error": "File must be a .docx file"}
            if not isinstance(level, int) or level < 1 or level > 9:
                return {"error": "Level must be an integer between 1 and 9"}
            doc = Document(filename)
            doc.add_heading(text, level=level)
            doc.save(filename)
            return {
                "success": True,
                "filename": filename,
                "text": text,
                "level": level,
                "message": f"Added heading '{text}' at level {level}"
            }
        except Exception as e:
            return {"error": f"Failed to add heading: {str(e)}"}
