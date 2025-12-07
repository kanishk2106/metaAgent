from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
def format_table(filename: str, table_index: int, has_header_row: bool = False, border_style: str = None, shading: list = None):
    doc = Document(filename)
    if table_index < len(doc.tables):
        table = doc.tables[table_index]
        if has_header_row and len(table.rows) > 0:
            table.style = 'Table Grid'
        if shading:
            for shade_info in shading:
                if isinstance(shade_info, dict):
                    row_idx = shade_info.get('row', 0)
                    color = shade_info.get('color', 'FFFFFF')
                    if row_idx < len(table.rows):
                        for cell in table.rows[row_idx].cells:
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:fill'), color)
                            cell._element.get_or_add_tcPr().append(shading_elm)
        doc.save(filename)
    return {"result": None}
