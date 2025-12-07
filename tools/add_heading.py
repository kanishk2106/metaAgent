from docx import Document
def add_heading(filename: str, text: str, level: int = 1):
    doc = Document(filename)
    doc.add_heading(text, level=level)
    doc.save(filename)
    return {"result": None}
