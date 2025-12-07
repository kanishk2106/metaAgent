from docx import Document
def add_table(filename: str, rows: int, cols: int, data: list = None):
    doc = Document(filename)
    table = doc.add_table(rows=rows, cols=cols)
    if data:
        for i, row_data in enumerate(data):
            if i < len(table.rows):
                row = table.rows[i]
                for j, cell_data in enumerate(row_data):
                    if j < len(row.cells):
                        row.cells[j].text = str(cell_data)
    doc.save(filename)
    return {"result": None}
