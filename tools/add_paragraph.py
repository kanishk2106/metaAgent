from docx import Document
def add_paragraph(filename: str, text: str, style: str = None):
    doc = Document(filename)
    paragraph = doc.add_paragraph(text)
    if style:
        paragraph.style = style
    doc.save(filename)
    return {"result": None}
