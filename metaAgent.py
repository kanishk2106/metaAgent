import os
import json
from pathlib import Path
from typing import Dict, Any, List
import google.generativeai as genai
FAMILIES_JSON = Path(
    "/Users/kanishk/Documents/Advanced NLP/Final_Project/MetaAgent/meta_families.json"
)
TOOLS_DIR = Path(
    "/Users/kanishk/Documents/Advanced NLP/Final_Project/MetaAgent/my_mcp_server/tools"
)
OUT_DIR = Path(
    "/Users/kanishk/Documents/Advanced NLP/Final_Project/MetaAgent/metaTools"
)
GEMINI_MODEL_NAME = "gemini-2.5-flash"
ADD_HEADING_EXAMPLE = '''"""
MCP Tool: add_heading
Add a heading to an existing Word document (.docx file).
"""
from docx import Document
import os
def register(server):
    @server.tool()
    def add_heading(filename: str, text: str, level: int = 1) -> dict:
        """
        Add a heading to a Word document.
        Args:
            filename: Path to the .docx file
            text: Heading text content
            level: Heading level (1-9, where 1 is largest)
        Returns:
            Success status and operation details
        """
        try:
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith('.docx'):
                return {"error": "File must be a .docx file"}
            if not isinstance(level, int) or level < 1 or level > 9:
                return {"error": "Level must be an integer between 1 and 9"}
            doc = Document(filename)
            doc.add_heading(text, level=level)
            doc.save(filename)
            return {
                "success": True,
                "filename": filename,
                "text": text,
                "level": level,
                "message": f"Added heading '{text}' at level {level}"
            }
        except Exception as e:
            return {"error": f"Failed to add heading: {str(e)}"}'''
ADD_PARAGRAPH_EXAMPLE = '''"""
MCP Tool: add_paragraph
Add a paragraph to an existing Word document (.docx file).
"""
from docx import Document
import os
def register(server):
    @server.tool()
    def add_paragraph(filename: str, text: str, style: str = None) -> dict:
        """
        Add a paragraph to a Word document.
        Args:
            filename: Path to the .docx file
            text: Paragraph text content
            style: Optional paragraph style name (e.g., 'Normal', 'BodyText')
        Returns:
            Success status and operation details
        """
        try:
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith('.docx'):
                return {"error": "File must be a .docx file"}
            doc = Document(filename)
            paragraph = doc.add_paragraph(text)
            if style:
                try:
                    paragraph.style = style
                except KeyError:
                    return {"error": f"Style '{style}' not found in document"}
            doc.save(filename)
            return {
                "success": True,
                "filename": filename,
                "text": text,
                "style": style if style else "default",
                "message": f"Added paragraph with {len(text)} characters"
            }
        except Exception as e:
            return {"error": f"Failed to add paragraph: {str(e)}"}'''
META_EXAMPLE_CODE = '''"""
MCP Tool: add_heading_with_optional_paragraph
Add a heading and optionally a paragraph to an existing Word document (.docx file).
"""
from docx import Document
import os
from typing import Optional, Dict, Any
def register(server):
    @server.tool()
    def add_heading_with_optional_paragraph(
        filename: str,
        heading_text: str,
        heading_level: int = 1,
        paragraph_text: Optional[str] = None,
        paragraph_style: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Add a heading and optionally a paragraph to a Word document.
        Mandatory behavior (FIRST TOOL: add_heading):
            - Always adds a heading.
        Optional behavior (SECOND TOOL: add_paragraph):
            - If `paragraph_text` is provided, also add a paragraph.
        Args:
            filename: Path to the .docx file.
            heading_text: Heading text content (required).
            heading_level: Heading level (1–9, where 1 is largest).
            paragraph_text: Optional paragraph text. If None, no paragraph is added.
            paragraph_style: Optional paragraph style (e.g., 'Normal', 'BodyText').
        Returns:
            A dict describing what operations were performed, or an error.
        """
        try:
            if not os.path.exists(filename):
                return {"error": f"File not found: {filename}"}
            if not filename.endswith(".docx"):
                return {"error": "File must be a .docx file"}
            if not isinstance(heading_level, int) or heading_level < 1 or heading_level > 9:
                return {"error": "heading_level must be an integer between 1 and 9"}
            doc = Document(filename)
            doc.add_heading(heading_text, level=heading_level)
            operations = {
                "heading_added": True,
                "heading_level": heading_level,
                "paragraph_added": False,
                "paragraph_style": None,
            }
            if paragraph_text is not None:
                paragraph = doc.add_paragraph(paragraph_text)
                if paragraph_style:
                    try:
                        paragraph.style = paragraph_style
                    except KeyError:
                        return {
                            "error": f"Style '{paragraph_style}' not found in document",
                            "filename": filename,
                        }
                operations["paragraph_added"] = True
                operations["paragraph_style"] = paragraph_style or "default"
            doc.save(filename)
            return {
                "success": True,
                "filename": filename,
                "heading_text": heading_text,
                "operations": operations,
                "message": (
                    "Added heading and paragraph"
                    if operations["paragraph_added"]
                    else "Added heading only"
                ),
            }
        except Exception as e:
            return {"error": f"Failed to add heading and optional paragraph: {str(e)}"}'''
def init_gemini() -> genai.GenerativeModel:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "Please set GEMINI_API_KEY environment variable, e.g.\n"
            "  export GEMINI_API_KEY='your-key-here'"
        )
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(GEMINI_MODEL_NAME)
def load_families() -> Dict[str, Any]:
    if not FAMILIES_JSON.exists():
        raise FileNotFoundError(f"Families JSON not found: {FAMILIES_JSON}")
    return json.loads(FAMILIES_JSON.read_text(encoding="utf-8"))
def load_tool_source(tool_name: str) -> str:
    path = TOOLS_DIR / f"{tool_name}.py"
    if not path.exists():
        return f"# WARNING: Source file for tool '{tool_name}' not found at {path}"
    return path.read_text(encoding="utf-8")
def clean_code(text: str) -> str:
    if not text:
        return text
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if not lines:
            return ""
        first = lines[0].strip()
        if first.startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).lstrip()
    return text
def build_prompt_for_family(family_name: str, cfg: Dict[str, Any]) -> str:
    first_tool = cfg["first_tool"]
    optional_tools: List[str] = cfg.get("all_optional_tools", [])
    tool_names: List[str] = [first_tool] + [t for t in optional_tools if t != first_tool]
    tool_sources: List[str] = []
    for name in tool_names:
        src = load_tool_source(name)
        tool_sources.append(f"# ===== Source for {name}.py =====\n{src}")
    tools_block = "\n\n".join(tool_sources)
    prompt = f"""
You are helping me generate HIGH-LEVEL "meta tools" for a local MCP-style tool server.
My tools are plain Python modules with this pattern:
- A top docstring starting with "MCP Tool: ..."
- Imports (docx, os, etc.) as needed.
- A `register(server)` function that defines one `@server.tool()` function.
- Each tool returns a JSON-serializable dict with either "success" or "error".
Below are TWO base tools and ONE combined meta tool. Treat these as the GOLD STANDARD
for style, structure, and the way a meta tool makes the FIRST TOOL mandatory and the
SECOND TOOL optional.
<<<BASE_TOOL_EXAMPLES_START>>>
{ADD_HEADING_EXAMPLE}
{ADD_PARAGRAPH_EXAMPLE}
<<<BASE_TOOL_EXAMPLES_END>>>
Now here is a COMBINED META TOOL built from those two base tools. This shows the
EXACT pattern I want: the first tool ("add_heading") is MANDATORY, and the second tool
("add_paragraph") is OPTIONAL and controlled by parameters.
<<<META_TOOL_EXAMPLE_START>>>
{META_EXAMPLE_CODE}
<<<META_TOOL_EXAMPLE_END>>>
================= FAMILY INFO: NEW META TOOL TO GENERATE =================
Family name (key): {family_name}
FIRST TOOL (MANDATORY): {first_tool}
OPTIONAL TOOLS: {optional_tools}
Here are the ACTUAL tool implementations for this family, loaded from the tool directory:
<<<FAMILY_TOOL_SOURCES_START>>>
{tools_block}
<<<FAMILY_TOOL_SOURCES_END>>>
Your job is to write a SINGLE META TOOL that:
  - ALWAYS executes the FIRST TOOL ({first_tool}) as a mandatory step.
  - CAN OPTIONALLY execute any subset of the OPTIONAL TOOLS {optional_tools},
    controlled by input parameters/flags.
You do NOT need to reproduce every possible sequence. Instead, design the meta tool's
arguments so they can express different combinations:
  - First tool always runs.
  - Each optional tool is enabled or disabled based on arguments.
=================================================================
REQUIREMENTS FOR THE NEW META TOOL YOU GENERATE:
1. FORMAT:
   - Use the SAME overall structure and coding style as the example meta tool:
     * A top docstring with "MCP Tool: <name>" and a concise description.
     * Imports as needed.
     * A `def register(server):` function.
     * Inside it, ONE `@server.tool()` function implementing the meta tool.
   - Return a JSON-serializable dict with either "success" or an "error" key.
2. NAMING:
   - Name the MCP tool and function based on the FIRST TOOL, for example:
       MCP Tool: {first_tool}_meta
       def {first_tool}_meta(...)
   - It is fine if you slightly adjust the name, but keep it clearly linked to {first_tool}.
3. SEMANTICS:
   - The FIRST TOOL ({first_tool}) is MANDATORY:
     * The meta tool must ALWAYS perform that behavior as the first step or core step.
   - The OTHER TOOLS {optional_tools} are OPTIONAL behaviors controlled by arguments.
     * Design parameters (e.g., booleans or optional values) so each optional behavior
       can be turned ON/OFF independently.
4. IMPLEMENTATION:
   - Implement the logic directly in this meta tool file, using the tool sources above as
     guidance for semantics and validation.
   - Reuse validation patterns from the base tools (e.g., file existence checks, extension checks).
   - Use appropriate libraries based on the tool implementations (e.g., docx, os, subprocess).
5. OUTPUT:
   - Return a dict summarizing the operations, for example:
     * any relevant path/filename
     * which optional behaviors actually ran
     * important outputs or results
     * an overall "message" and a "success" flag.
   - For any error, return `{{"error": "...message..."}}` (do not raise).
6. IMPORTANT:
   - Output ONLY the complete Python module code for the NEW meta tool.
   - DO NOT wrap the output in backticks.
   - DO NOT include any explanation or commentary outside the code.
"""
    return prompt.strip()
def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    families = load_families()
    model = init_gemini()
    print(f"[INFO] Loaded {len(families)} families from {FAMILIES_JSON}")
    for family_name, cfg in families.items():
        if "first_tool" not in cfg or "all_optional_tools" not in cfg:
            print(f"[WARN] Skipping family {family_name}: missing first_tool or all_optional_tools")
            continue
        print(f"\n[INFO] Generating meta tool for family: {family_name}")
        prompt = build_prompt_for_family(family_name, cfg)
        resp = model.generate_content(prompt)
        raw_text = getattr(resp, "text", "") or ""
        code = clean_code(raw_text)
        if not code:
            print(f"[WARN] Empty response for family {family_name}, skipping.")
            if raw_text:
                print(f"[DEBUG] Raw response snippet: {raw_text[:200]!r}")
            continue
        safe_name = "".join(
            c if c.isalnum() or c in ("_", "-") else "_" for c in family_name
        )
        out_path = OUT_DIR / f"meta_family_{safe_name}.py"
        out_path.write_text(code, encoding="utf-8")
        print(f"[INFO] Wrote {out_path}")
if __name__ == "__main__":
    main()
